"""Generate Crystal's AI Teaching Institute Coach application letter."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Header
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = header.add_run('Crystal Gutierrez')
r.bold = True
r.font.size = Pt(12)
header.add_run('\nAdjunct Professor, New Mexico State University')
header.add_run('\nM.S. Candidate, Teaching and Curriculum Building, NMSU')
header.add_run('\ncag1145@nmsu.edu')

doc.add_paragraph()

# ============================================================
# SECTION 1
# ============================================================
h1 = doc.add_paragraph()
r = h1.add_run('AI Experience or AI-Augmented Teaching Experience and/or Training')
r.bold = True
r.font.size = Pt(11)

doc.add_paragraph(
    'I work across both AI research and classroom instruction. As a researcher '
    'at Axion Deep Labs, I co-authored a study applying persistent homology to '
    'neural network loss landscapes for predicting catastrophic forgetting, '
    'currently under submission to ArXiv. My daily toolkit includes Python, '
    'PyTorch, NumPy, Pandas, Jupyter, R, XGBoost, and various regression methods.'
)

doc.add_paragraph(
    'I conducted applied machine learning research in collaboration with Bayer '
    'and the Purdue Data Mine, building and evaluating ML models on real-world '
    'agricultural and industrial datasets.'
)

doc.add_paragraph(
    'As an adjunct professor at NMSU, I incorporate AI tools directly into my '
    'teaching. Students use large language models as supplementary tutoring tools, '
    'critically evaluate AI-generated solutions against their own work, and learn '
    'to use Jupyter notebooks and Python libraries for data analysis. I emphasize '
    'AI literacy as a core competency: not just how to use AI tools, but how to '
    'assess their outputs for correctness and appropriateness.'
)

doc.add_paragraph(
    'My masters work in teaching and curriculum building at NMSU gives me formal '
    'training in pedagogical design, which I apply directly to structuring how '
    'students and colleagues learn to work with AI.'
)

# ============================================================
# SECTION 2
# ============================================================
h2 = doc.add_paragraph()
r = h2.add_run('Experience and Approach with/to Coaching')
r.bold = True
r.font.size = Pt(11)

doc.add_paragraph(
    'As an adjunct professor and a graduate student in teaching and curriculum '
    'building, I regularly mentor students and collaborate with faculty on course '
    'design. At Axion Deep Labs, I help translate complex ML concepts into '
    'accessible language for team members at different skill levels, guiding them '
    'through experimental design, statistical analysis, and technical writing.'
)

doc.add_paragraph(
    'Three strategies I have found effective when coaching faculty on AI adoption:'
)

p1 = doc.add_paragraph()
r = p1.add_run('Start with their discipline. ')
r.bold = True
p1.add_run(
    'Faculty adopt AI tools when they see immediate relevance. I begin by asking '
    'what problems they face in teaching and identify where AI can address those '
    'specific pain points.'
)

p2 = doc.add_paragraph()
r = p2.add_run('Demonstrate before prescribing. ')
r.bold = True
p2.add_run(
    'Live demonstrations using a colleague\'s actual course material are far more '
    'persuasive than documentation. Seeing an AI tool generate a rubric from their '
    'own assignment narrows the gap between skepticism and adoption immediately.'
)

p3 = doc.add_paragraph()
r = p3.add_run('Build judgment, not just tool proficiency. ')
r.bold = True
p3.add_run(
    'The most important coaching outcome is that faculty develop the judgment to '
    'evaluate AI outputs, set appropriate boundaries for student use, and design '
    'assessments that remain meaningful in an AI-augmented environment.'
)

doc.add_paragraph(
    'My combination of active AI research, formal training in curriculum design, '
    'and classroom teaching positions me to help colleagues navigate AI integration '
    'with both technical depth and pedagogical grounding.'
)

# Save
output_path = '/home/joshua/projects/axiondeep-research/drafts/crystal_ai_coach_application.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
